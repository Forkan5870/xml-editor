import docx


def CreateDocx(array):
    
    doc = docx.Document()

    for i in range(len(array)):
        nombre_para = doc.add_paragraph()
        info_para = doc.add_paragraph()
        extra_para = doc.add_paragraph()
        ensayo1_para = doc.add_paragraph()
        ensayo2_para = doc.add_paragraph()
        comentarios_para = doc.add_paragraph()

        run = nombre_para.add_run(array[i]["nombre"])
        run.font.name = 'Verdana'
        run.bold = True
        run.font.size = docx.shared.Pt(25)

        info_para.style.font.name = 'Assistant'
        info_para.add_run("Instituto: ").bold = True
        info_para.add_run(f"{array[i]['instituto']}, ")
        info_para.add_run("Provincia: ").bold = True
        info_para.add_run(f"{array[i]['provincia']}, ")
        info_para.add_run("ESO: ").bold = True
        info_para.add_run(f"{array[i]['eso']}, ")
        info_para.add_run("Bachillerato: ").bold = True
        info_para.add_run(f"{array[i]['bachillerato']}, ")
        info_para.add_run("Inglés: ").bold = True
        info_para.add_run(f"{array[i]['ingles']}, ")
        info_para.add_run("Grado: ").bold = True
        info_para.add_run(array[i]['grado'])

        extra_para.style.font.name = 'Assistant'
        extra_para.add_run("Extraescolares: ").bold = True
        extra_para.add_run(array[i]['extraescolares'])

        ensayo1_para.style.font.name = 'Assistant'
        ensayo1_para.add_run("Ensayo 1: ").bold = True
        ensayo1_para.add_run(array[i]['ensayo1'])

        ensayo2_para.style.font.name = 'Assistant'
        ensayo2_para.add_run("Ensayo 2: ").bold = True
        ensayo2_para.add_run(array[i]['ensayo2'])

        comentarios_para.style.font.name = 'Assistant'
        comentarios_para.add_run("Comentarios: ").bold = True
        comentarios_para.add_run(array[i]['comentarios'])

        if (not(i == len(array) - 1)):
            doc.add_page_break() 

    doc.save('files/test.docx') 